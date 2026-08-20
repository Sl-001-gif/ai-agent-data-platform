# -*- coding: utf-8 -*-
"""Markdown -> Word 转换脚本（实习报告专用）。

用法：
  python tools/实习报告转word.py                      # 默认转换 v3
  python tools/实习报告转word.py 输入.md 输出.docx

规则：红色文字 = 待填写提示；## 一级标题、### 二级标题、#### 三级标题；
表格用 Markdown 管道表格；--- 为分页；> 引用为提示条；
<!-- COVER-START --> ... <!-- COVER-END --> 为封面下划线信息栏。
"""
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH = sys.argv[1] if len(sys.argv) > 1 else os.path.join(ROOT, "docs", "实习报告-数据分析实习生-v3.md")
DOCX_PATH = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(MD_PATH)[0] + ".docx"

RED = RGBColor(0xC0, 0x00, 0x00)
GRAY = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0, 0, 0)
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def _set_font(run, ea="宋体", ascii_font="Times New Roman", size=12, bold=None, color=None):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), ea)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_inline(paragraph, text, size=12, base_bold=False, color=None):
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _set_font(paragraph.add_run(part[2:-2]), size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            _set_font(paragraph.add_run(part[1:-1]), ea="宋体", ascii_font="Consolas", size=size - 0.5, bold=base_bold, color=color)
        else:
            _set_font(paragraph.add_run(part), size=size, bold=base_bold, color=color)


def _style_heading(doc, name, size):
    st = doc.styles[name]
    st.font.name = "黑体"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = BLACK
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "黑体")


def is_table_row(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def _add_cover_form(rows, doc):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblpr.append(borders)
    for r, (label, value) in enumerate(rows):
        c0, c1 = table.rows[r].cells
        c0.width, c1.width = Cm(4.2), Cm(8.0)
        tr = table.rows[r]._tr
        trpr = tr.get_or_add_trPr()
        th = OxmlElement("w:trHeight")
        th.set(qn("w:val"), "480")
        th.set(qn("w:hRule"), "atLeast")
        trpr.append(th)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p0.add_run(label), size=12)
        p1 = c1.paragraphs[0]
        if value.startswith("【"):
            _set_font(p1.add_run(value), size=12, color=RED)
        else:
            _set_font(p1.add_run(value), size=12)
        tcpr = c1._tc.get_or_add_tcPr()
        tcb = OxmlElement("w:tcBorders")
        btm = OxmlElement("w:bottom")
        btm.set(qn("w:val"), "single")
        btm.set(qn("w:sz"), "8")
        btm.set(qn("w:color"), "000000")
        tcb.append(btm)
        tcpr.append(tcb)


def build_doc(md_text, doc):
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    _style_heading(doc, "Heading 1", 16)
    _style_heading(doc, "Heading 2", 14)
    _style_heading(doc, "Heading 3", 12)
    for sec in doc.sections:
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
        sec.top_margin, sec.bottom_margin = Cm(2.54), Cm(2.54)
        sec.left_margin, sec.right_margin = Cm(3.0), Cm(2.6)

    lines = md_text.splitlines()
    i = 0
    cover_done = False
    while i < len(lines):
        stripped = lines[i].strip()

        if not cover_done and re.sub(r"^#+\s*", "", stripped) == "华北水利水电大学":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(72)
            p.paragraph_format.space_after = Pt(24)
            _set_font(p.add_run(re.sub(r"^#+\s*", "", stripped)), ea="黑体", ascii_font="黑体", size=28, bold=True)
            i += 1
            continue
        if not cover_done and re.sub(r"^#+\s*", "", stripped) == "本科生实习报告":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(18)
            _set_font(p.add_run(re.sub(r"^#+\s*", "", stripped)), ea="黑体", ascii_font="黑体", size=22, bold=True)
            i += 1
            continue
        if not cover_done and stripped.startswith("（如不想突出 AI"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(18)
            add_inline(p, stripped, size=10.5, color=GRAY)
            i += 1
            continue

        if stripped == "<!-- COVER-START -->":
            rows = []
            i += 1
            while i < len(lines) and lines[i].strip() != "<!-- COVER-END -->":
                s = lines[i].strip()
                if s and not s.startswith("#"):
                    if s.startswith("报告题目："):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(30)
                        p.paragraph_format.space_after = Pt(6)
                        _set_font(p.add_run("报告题目："), ea="黑体", ascii_font="黑体", size=14)
                        _set_font(p.add_run(s[len("报告题目："):]), ea="黑体", ascii_font="黑体", size=16, bold=True)
                    else:
                        label, _, value = s.partition("：")
                        rows.append((label, value))
                i += 1
            i += 1
            if rows:
                _add_cover_form(rows, doc)
            continue

        if is_table_row(stripped):
            rows = []
            while i < len(lines) and is_table_row(lines[i]):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r, row in enumerate(rows):
                    for c, val in enumerate(row):
                        cell = table.cell(r, c)
                        cell.width = Cm(4.5 if c == 0 else 10.5)
                        para = cell.paragraphs[0]
                        if r == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            _set_font(para.add_run(val), ea="黑体", ascii_font="黑体", size=11, bold=True)
                            tcpr = cell._element.get_or_add_tcPr()
                            shd = tcpr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "D9D9D9"})
                            tcpr.append(shd)
                        elif "请填写" in val or "请补充" in val or val.startswith("【"):
                            _set_font(para.add_run(val), size=11, color=RED)
                        else:
                            _set_font(para.add_run(val), size=11)
            continue

        if stripped in ("---", "***"):
            doc.add_page_break()
            cover_done = True
            i += 1
            continue

        if stripped.startswith("#### "):
            p = doc.add_heading(level=3)
            add_inline(p, stripped[5:], size=12, base_bold=True)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(level=2)
            add_inline(p, stripped[4:], size=14, base_bold=True)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(level=1)
            add_inline(p, stripped[3:], size=16, base_bold=True)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(level=1)
            add_inline(p, stripped[2:], size=16, base_bold=True)
            i += 1
            continue

        if re.match(r"^\s*[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^\s*[-*]\s+", "", stripped), size=12)
            p.paragraph_format.line_spacing = 1.5
            i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.line_spacing = 1.5
            add_inline(p, stripped, size=12)
            i += 1
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, stripped.lstrip("> "), size=10.5, color=RED, base_bold=True)
            i += 1
            continue
        if stripped.startswith("【✍") or stripped.startswith("（插入"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, stripped, size=11, color=RED, base_bold=True)
            i += 1
            continue
        if re.match(r"^\d{4}\s*年\s*\d{1,2}\s*月$", stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            _set_font(p.add_run(stripped), size=14)
            i += 1
            continue

        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.line_spacing = 1.5
            add_inline(p, stripped, size=12)
        i += 1


def main():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md = f.read()
    doc = Document()
    build_doc(md, doc)
    doc.save(DOCX_PATH)
    print("OK:", DOCX_PATH)


if __name__ == "__main__":
    main()
